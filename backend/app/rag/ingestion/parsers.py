"""Document parsers: PDF, DOCX, and scanned images (OCR).

Each parser implements the ``DocumentParser`` protocol and is chosen by file extension via
``ParserRegistry`` — adding a new format means adding a parser, not editing existing code
(Open/Closed principle). Parsing produces located ``ParsedElement``s (text/heading/table with
page + bbox) that feed the chunking pipeline.
"""

from __future__ import annotations

import statistics
from pathlib import Path
from typing import Protocol

from app.rag.chunking.models import ElementType, ParsedDocument, ParsedElement


class ParserError(RuntimeError):
    pass


class DocumentParser(Protocol):
    extensions: tuple[str, ...]

    def parse(self, path: str, *, title: str) -> ParsedDocument: ...


def _bbox(coords) -> dict | None:  # noqa: ANN001
    if not coords:
        return None
    x0, y0, x1, y1 = coords
    return {"x0": float(x0), "y0": float(y0), "x1": float(x1), "y1": float(y1)}


def _rows_to_markdown(rows: list[list]) -> str:
    cleaned = [[("" if c is None else str(c)).strip() for c in row] for row in rows if row]
    if not cleaned:
        return ""
    width = max(len(r) for r in cleaned)
    cleaned = [r + [""] * (width - len(r)) for r in cleaned]
    header = "| " + " | ".join(cleaned[0]) + " |"
    sep = "| " + " | ".join("---" for _ in range(width)) + " |"
    body = ["| " + " | ".join(r) + " |" for r in cleaned[1:]]
    return "\n".join([header, sep, *body])


class PdfParser:
    extensions = (".pdf",)

    def parse(self, path: str, *, title: str) -> ParsedDocument:
        import fitz  # PyMuPDF

        elements: list[ParsedElement] = []
        doc = fitz.open(path)
        try:
            for page_no, page in enumerate(doc, start=1):
                elements.extend(self._page_blocks(page, page_no))
            page_count = doc.page_count
        finally:
            doc.close()

        elements.extend(self._tables(path))
        return ParsedDocument(
            title=title, file_type="pdf", elements=elements, page_count=page_count
        )

    @staticmethod
    def _page_blocks(page, page_no: int) -> list[ParsedElement]:  # noqa: ANN001
        data = page.get_text("dict")
        spans = [
            span
            for block in data.get("blocks", [])
            for line in block.get("lines", [])
            for span in line.get("spans", [])
        ]
        sizes = [s["size"] for s in spans if s.get("text", "").strip()]
        median = statistics.median(sizes) if sizes else 0.0

        out: list[ParsedElement] = []
        for block in data.get("blocks", []):
            lines = block.get("lines", [])
            if not lines:
                continue
            text = " ".join(span["text"] for line in lines for span in line["spans"]).strip()
            if not text:
                continue
            block_size = max(
                (span["size"] for line in lines for span in line["spans"]), default=median
            )
            is_heading = median and block_size > median * 1.25 and len(text) < 80
            out.append(
                ParsedElement(
                    text=text,
                    type=ElementType.HEADING if is_heading else ElementType.TEXT,
                    page=page_no,
                    bbox=_bbox(block.get("bbox")),
                )
            )
        return out

    @staticmethod
    def _tables(path: str) -> list[ParsedElement]:
        try:
            import pdfplumber
        except ImportError:  # pragma: no cover
            return []
        out: list[ParsedElement] = []
        try:
            with pdfplumber.open(path) as pdf:
                for page_no, page in enumerate(pdf.pages, start=1):
                    for table in page.find_tables():
                        markdown = _rows_to_markdown(table.extract())
                        if markdown:
                            out.append(
                                ParsedElement(
                                    text=markdown,
                                    type=ElementType.TABLE,
                                    page=page_no,
                                    bbox=_bbox(table.bbox),
                                )
                            )
        except Exception:  # noqa: BLE001 - table extraction is best-effort
            return out
        return out


class DocxParser:
    extensions = (".docx",)

    def parse(self, path: str, *, title: str) -> ParsedDocument:
        import docx

        document = docx.Document(path)
        elements: list[ParsedElement] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            etype = ElementType.HEADING if style.startswith("heading") else ElementType.TEXT
            elements.append(ParsedElement(text=text, type=etype, page=None))

        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            markdown = _rows_to_markdown(rows)
            if markdown:
                elements.append(ParsedElement(text=markdown, type=ElementType.TABLE, page=None))

        return ParsedDocument(title=title, file_type="docx", elements=elements)


class ImageParser:
    extensions = (".png", ".jpg", ".jpeg", ".tiff", ".bmp")

    def parse(self, path: str, *, title: str) -> ParsedDocument:
        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:  # pragma: no cover
            raise ParserError("OCR dependencies not available") from exc
        try:
            text = pytesseract.image_to_string(Image.open(path))
        except Exception as exc:  # noqa: BLE001
            raise ParserError(
                "OCR failed — is the Tesseract binary installed and on PATH?"
            ) from exc
        text = text.strip()
        elements = [ParsedElement(text=text, type=ElementType.TEXT, page=1)] if text else []
        return ParsedDocument(title=title, file_type="image", elements=elements, page_count=1)


class ParserRegistry:
    def __init__(self, parsers: list[DocumentParser] | None = None) -> None:
        self._parsers = parsers or [PdfParser(), DocxParser(), ImageParser()]

    def for_extension(self, ext: str) -> DocumentParser:
        ext = ext.lower()
        for parser in self._parsers:
            if ext in parser.extensions:
                return parser
        raise ParserError(f"Unsupported file type: {ext}")

    def parse(self, path: str, *, title: str | None = None) -> ParsedDocument:
        ext = Path(path).suffix
        return self.for_extension(ext).parse(path, title=title or Path(path).name)
